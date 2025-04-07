# src/pubmed_fetcher/writer.py

from docx import Document
import os

def save_to_docx(pmid, metadata, html_text, pdf_text, output_dir):
    doc = Document()
    doc.add_heading(metadata.get("title", "No Title"), 0)

    doc.add_paragraph(f"PMID: {pmid}")
    doc.add_paragraph(f"DOI: {metadata.get('doi', 'N/A')}")
    doc.add_paragraph(f"Authors: {metadata.get('authors', 'N/A')}")

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(metadata.get("abstract", "No Abstract Available"))

    doc.add_heading("HTML Full Text", level=1)
    doc.add_paragraph(html_text or "No HTML text available")

    doc.add_heading("PDF Full Text", level=1)
    doc.add_paragraph(pdf_text or "No PDF text available")

    os.makedirs(output_dir, exist_ok=True)
    filepath = os.path.join(output_dir, f"{pmid}.docx")
    doc.save(filepath)
    print(f"Saved: {filepath}")
